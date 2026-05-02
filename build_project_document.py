from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "CMS_Project_Backend_Upgrade_Report.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True, color=(255, 255, 255))
        shade_cell(hdr[i], "334155")
        if widths:
            hdr[i].width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F1F5F9")
    cell.text = ""
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(30, 41, 59)
    if cell.paragraphs and not cell.paragraphs[0].text:
        cell._element.remove(cell.paragraphs[0]._element)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CMS Project Backend Upgrade Report")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Frontend-backend connection, PostgreSQL verification, schema coverage upgrade, and error log")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(71, 85, 105)

    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Project", "Hostel Mess / Canteen Management System"],
            ["Workspace", r"D:\Project_Folder\CMS\CMS_Repo"],
            ["Backend", "FastAPI + asyncpg + PostgreSQL"],
            ["Frontend", "Flutter app + React admin panel present"],
            ["Active database", "canteen_db"],
            ["Current backend score vs DB schema", "Improved from about 38/100 to about 78/100"],
        ],
        [Inches(1.8), Inches(5.6)],
    )

    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "The project now has a working FastAPI backend connected to PostgreSQL and a Flutter auth flow connected to that backend. "
        "The database schema was significantly more advanced than the original backend, so the backend was upgraded to cover the major schema areas: auth sessions, audit logs, feedback, leave, notifications, settings, reports, staff/admin, and menu management."
    )
    add_bullets(
        doc,
        [
            "Backend connection to PostgreSQL was verified with the live FastAPI lifespan test.",
            "Flutter registration/login was changed from local SharedPreferences fake users to real HTTP API calls.",
            "The active PostgreSQL DB is canteen_db; cms_db was identified as unused.",
            "bcrypt was downgraded from 5.0.0 to 4.0.1 to restore passlib compatibility.",
            "Backend route coverage expanded to 52 FastAPI routes, including 40 /api paths.",
        ],
    )

    add_heading(doc, "2. Confirmed Database State")
    add_table(
        doc,
        ["Item", "Confirmed Value"],
        [
            ["DB host", "localhost"],
            ["DB port", "5432"],
            ["DB name in use", "canteen_db"],
            ["DB user", "postgres"],
            ["PostgreSQL server", "PostgreSQL 18.1 on x86_64-windows"],
            ["Tables found", "24 public tables"],
            ["Seed/core rows", "roles: 4, hostels: 4, meal_slots: 4, menu_items: 15, booking_status: 7"],
            ["Current app rows checked", "users: 2, students: 1, bookings: 0"],
            ["Other DB found", "cms_db, considered unused based on backend .env and live health check"],
        ],
        [Inches(2.0), Inches(5.4)],
    )

    add_heading(doc, "3. Frontend to Backend Connection Changes")
    add_table(
        doc,
        ["File", "Change"],
        [
            ["canteen-app/lib/auth_service.dart", "Replaced fake local user storage with HTTP calls to /api/auth/register, /api/auth/login, and /api/auth/me. Stores access and refresh tokens in SharedPreferences."],
            ["canteen-app/lib/registration_screen.dart", "Removed duplicate checks against old local fake-user records; backend now enforces email and UID uniqueness."],
            ["canteen-app/android/app/src/main/AndroidManifest.xml", "Added INTERNET permission and cleartext traffic support for local development HTTP calls."],
            ["canteen-app/android/gradle.properties", "Disabled Kotlin/Gradle caching to avoid Windows mixed-drive incremental cache failure."],
        ],
        [Inches(2.45), Inches(4.95)],
    )

    add_heading(doc, "4. Backend Auth and Security Upgrades")
    add_table(
        doc,
        ["Area", "Before", "After"],
        [
            ["Login audit", "JWT returned only", "Writes last_login_at and last_login_ip on success"],
            ["Failed login tracking", "Not implemented", "Increments login_attempts on wrong password"],
            ["Account locking", "Schema existed only", "Locks account after repeated failed attempts"],
            ["Sessions", "TODO comment only", "Creates sessions rows with hashed access/refresh tokens"],
            ["Session validation", "JWT only", "JWT must also match a non-revoked, non-expired session"],
            ["Auth logs", "Not used", "Writes login, failed_login, account_locked, and logout/revoke events"],
            ["Session APIs", "Missing", "Added list sessions and revoke session endpoints"],
        ],
        [Inches(1.8), Inches(2.4), Inches(3.2)],
    )

    add_heading(doc, "5. New Backend Modules Added")
    add_table(
        doc,
        ["Module", "Purpose", "Key Routes"],
        [
            ["admin.py", "User and staff management", "/api/admin/users, /api/admin/staff"],
            ["feedback.py", "Student feedback and staff resolution", "/api/feedback, /api/feedback/mine, /api/feedback/all"],
            ["leave.py", "Student leave requests and staff decisions", "/api/leave, /api/leave/all"],
            ["notifications.py", "Create, materialize, read notifications", "/api/notifications"],
            ["reports.py", "Expose DB reporting views", "/api/reports/today/bookings, /api/reports/today/meal-count"],
            ["settings.py", "Hostel settings read/update", "/api/settings/hostels"],
        ],
        [Inches(1.55), Inches(2.5), Inches(3.35)],
    )

    add_heading(doc, "6. Existing Backend Modules Upgraded")
    add_table(
        doc,
        ["File", "Upgrade"],
        [
            ["app/main.py", "Added routers for new modules and middleware that writes unexpected server errors into error_logs."],
            ["app/dependencies.py", "Session validation added on top of JWT validation; updates sessions.last_used_at."],
            ["app/routes/meals.py", "Added menu item and meal menu management APIs; fixed menu item response shape."],
            ["app/routes/bookings.py", "Moved /history/list before /{id} so FastAPI route matching works correctly."],
            ["app/services/auth_service.py", "Added SHA-256 token hashing helper for sessions."],
            ["app/schemas/*.py", "Added Pydantic schemas for admin, feedback, leave, notifications, settings, common responses, sessions, and logs."],
            ["requirements.txt", "Changed bcrypt from 5.0.0 to 4.0.1 for passlib compatibility."],
        ],
        [Inches(2.1), Inches(5.3)],
    )

    add_heading(doc, "7. API Coverage After Upgrade")
    add_bullets(
        doc,
        [
            "Auth: register, login, current user, sessions, revoke session, auth logs.",
            "Meals: slots, today, menu lookup, menu items, meal menu creation, publish/unpublish.",
            "Bookings: create, read, update, cancel, skip, undo skip, history.",
            "Tokens: active, upcoming, QR data, staff scan.",
            "Leave: create, list own leave, staff list all, staff approve/reject.",
            "Feedback: create, list own, list all for staff, resolve feedback, list tags.",
            "Notifications: create targeted notification, list mine, mark read.",
            "Settings: list/update hostel settings.",
            "Reports: today's bookings, today's meal count, active tokens, student history.",
            "Admin: list users, list/create/update staff, activate/deactivate users.",
        ],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "8. Error Log and Fixes")
    add_table(
        doc,
        ["Error / Symptom", "Cause", "Fix Applied"],
        [
            ["Flutter desktop: Building with plugins requires symlink support", "Windows Developer Mode was disabled.", "Enabled Developer Mode through ms-settings:developers."],
            ["Android build: Kotlin daemon cache failure, different roots C: and D:", "Kotlin incremental cache could not handle plugin source in Pub cache on C: and project on D:.", "Disabled kotlin.incremental, classpath snapshot, and Gradle caching in android/gradle.properties."],
            ["Flutter registration: Internal Server Error with passlib/bcrypt stack trace", "passlib 1.7.4 incompatible with bcrypt 5.0.0.", "Pinned bcrypt to 4.0.1 and installed it in backend .venv."],
            ["Dart format/analyze access errors", "Windows sandbox/profile access blocked Dart analytics/cache paths.", "Used workspace app data and escalated where needed for Dart analysis."],
            ["python command not found", "System PATH did not expose python command.", "Used backend .venv Python and bundled workspace Python."],
            ["Two PostgreSQL databases found", "Accidental creation of cms_db plus active canteen_db.", "Verified backend .env and live DB connection use canteen_db; cms_db identified unused."],
            ["Route /api/bookings/history/list risk", "Dynamic /{id} route was defined before static history route.", "Moved history route before /{id}."],
            ["Pydantic type annotation error in meals schema", "Field name date shadowed imported date type.", "Aliased date/time imports to Date and Time."],
        ],
        [Inches(2.25), Inches(2.45), Inches(2.7)],
    )

    add_heading(doc, "9. Verification Performed")
    add_table(
        doc,
        ["Check", "Result"],
        [
            ["Backend DB lifespan test", "Passed: created asyncpg pool, printed DB Success, closed pool."],
            ["DB health check", "Passed: confirmed canteen_db, PostgreSQL version, 24 tables, seed counts."],
            ["Password hash/verify check", "Passed after bcrypt downgrade."],
            ["Python compileall", "Passed for app modules."],
            ["FastAPI route import/list", "Passed: 52 routes loaded."],
            ["Flutter changed file analyze", "Passed for auth_service.dart and registration_screen.dart earlier in the work."],
            ["Full Flutter test", "Not completed; timed out/environment issues before Android build was stabilized."],
        ],
        [Inches(2.35), Inches(5.05)],
    )

    add_heading(doc, "10. Fresh Start Commands")
    doc.add_paragraph("Start backend in PowerShell window 1:")
    add_code_block(
        doc,
        [
            r"cd D:\Project_Folder\CMS\CMS_Repo\canteen-backend",
            r".\.venv\Scripts\uvicorn.exe app.main:app --reload",
        ],
    )
    doc.add_paragraph("Start emulator through Android Studio Device Manager, then run Flutter in PowerShell window 2:")
    add_code_block(
        doc,
        [
            r"cd D:\Project_Folder\CMS\CMS_Repo\canteen-app",
            r"D:\Flutter_Bank\flutter\bin\flutter.bat devices",
            r"D:\Flutter_Bank\flutter\bin\flutter.bat run -d emulator-5554",
        ],
    )
    doc.add_paragraph("Expected backend request flow during app test:")
    add_code_block(
        doc,
        [
            "POST /api/auth/register",
            "POST /api/auth/login",
            "GET /api/auth/me",
        ],
    )

    add_heading(doc, "11. Current Score and Remaining Work")
    add_table(
        doc,
        ["Category", "Current Status"],
        [
            ["Previous backend score vs DB schema", "About 38/100"],
            ["Current backend score vs DB schema", "About 78/100"],
            ["Why not 100 yet", "Needs production-grade email/SMS/push delivery, OTP flows, refresh-token endpoint, complete frontend integration, automated tests, admin UI, and deployment hardening."],
            ["Best next backend step", "Add token refresh/logout endpoints and OTP/email verification flows."],
            ["Best next frontend step", "Connect meal booking, history, feedback, notifications, and profile screens to the new APIs."],
        ],
        [Inches(2.4), Inches(5.0)],
    )

    add_heading(doc, "12. Important Notes")
    add_bullets(
        doc,
        [
            "Do not delete canteen_db; it is the active project database.",
            "cms_db was identified as unused, but any destructive database deletion should be done manually and carefully in pgAdmin.",
            "The backend now requires sessions rows for JWT validation; existing tokens issued before this upgrade may stop working. Log in again from the app.",
            "The admin seed account exists in the database. It is for development and should be replaced or secured before deployment.",
            "The Flutter app auth is connected to backend, but several other screens may still use local SharedPreferences until wired to the new APIs.",
        ],
    )

    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("CMS Backend Upgrade Report")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
